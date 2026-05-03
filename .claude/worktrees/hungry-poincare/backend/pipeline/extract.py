"""
Multi-format içerik çıkarıcılar.

Desteklenen formatlar:
  - HTML  → BeautifulSoup (gürültü temizleme + tablo → Markdown)
  - PDF   → pdfplumber (birincil) / pypdf (yedek)
  - DOCX  → python-docx
  - MD    → Doğrudan metin

Çıktı: RawDocument
"""
from __future__ import annotations

import io
import re
from dataclasses import dataclass, field
from urllib.parse import urlparse

from bs4 import BeautifulSoup


# ---------------------------------------------------------------------------
# Veri Modeli
# ---------------------------------------------------------------------------

@dataclass
class RawDocument:
    url: str
    title: str
    body: str                        # Yapılandırılmış metin (## başlıklar, tablolar)
    fmt: str                         # html / pdf / docx / md
    metadata: dict = field(default_factory=dict)

    def is_empty(self) -> bool:
        return not self.body.strip()


# ---------------------------------------------------------------------------
# HTML Çıkarıcı
# ---------------------------------------------------------------------------

_HTML_NOISE_TAGS = frozenset({
    "script", "style", "noscript", "iframe", "link",
    "meta", "button", "input", "select", "textarea",
})

_HTML_NOISE_KEYWORDS = frozenset({
    "nav", "navbar", "navigation", "menu", "header",
    "footer", "sidebar", "breadcrumb", "pagination",
    "social", "share", "cookie", "gdpr", "banner",
    "advertisement", "popup", "modal", "overlay",
    "topbar", "menubar", "dropdown",
})

_HTML_MAIN_SELECTORS = [
    "article", "main", "[role='main']",
    "#content", "#main-content", "#page-content", "#main",
    ".content", ".main-content", ".page-content",
    ".entry-content", ".post-content", ".article-body",
]

_HTML_BLOCK_CONTENT = frozenset({
    "h1", "h2", "h3", "h4", "h5", "h6",
    "p", "li", "pre", "blockquote", "figcaption", "dt", "dd",
})


class HTMLExtractor:
    def extract(self, html: str, url: str) -> RawDocument | None:
        soup = BeautifulSoup(html, "lxml")

        title = self._get_title(soup)

        # Gürültü etiketlerini kaldır
        for tag in soup(list(_HTML_NOISE_TAGS)):
            tag.decompose()

        # id/class tabanlı gürültü kaldırma
        # el.parent is None  →  parent zaten decompose edildi, bu child'ı atla
        # (BS4, parent.decompose() sırasında child.attrs referansını temizler;
        #  sonradan el.get() çağrısı AttributeError verir)
        for el in soup.find_all(True):
            if el.parent is None:
                continue
            if self._is_noise_element(el):
                el.decompose()

        # Ana içerik alanını bul
        main = None
        for sel in _HTML_MAIN_SELECTORS:
            main = soup.select_one(sel)
            if main:
                break
        if main is None:
            main = soup.find("body") or soup

        body = self._extract_text(main)

        if not body.strip():
            return None

        source = urlparse(url).netloc or "belek.edu.tr"
        return RawDocument(
            url=url,
            title=title,
            body=body,
            fmt="html",
            metadata={"source": source},
        )

    def _is_noise_element(self, el) -> bool:
        try:
            el_id      = str(el.get("id")    or "").lower()
            el_classes = " ".join(el.get("class") or []).lower()
            combined   = f"{el_id} {el_classes}"
            return any(kw in combined for kw in _HTML_NOISE_KEYWORDS)
        except Exception:
            return False

    def _get_title(self, soup: BeautifulSoup) -> str:
        og = soup.find("meta", property="og:title")
        if og and og.get("content"):
            return og["content"].strip()
        h1 = soup.find("h1")
        if h1:
            return h1.get_text(strip=True)
        tag = soup.find("title")
        if tag:
            return tag.get_text(strip=True).split("|")[0].split("-")[0].strip()
        return ""

    def _extract_text(self, root) -> str:
        parts: list[str] = []

        for el in root.find_all(list(_HTML_BLOCK_CONTENT) + ["table"], recursive=True):
            # İç içe geçme: üst eleman zaten işlendi mi?
            if el.parent and el.parent.name in _HTML_BLOCK_CONTENT:
                continue

            name = el.name

            if name in ("h1", "h2"):
                text = el.get_text(strip=True)
                if text:
                    parts.append(f"\n## {text}\n")

            elif name in ("h3", "h4", "h5", "h6"):
                text = el.get_text(strip=True)
                if text:
                    parts.append(f"\n### {text}\n")

            elif name == "table":
                md = _table_to_markdown(el)
                if md:
                    parts.append(f"\n{md}\n")

            else:  # p, li, pre, blockquote, dt, dd, figcaption
                text = el.get_text(separator=" ", strip=True)
                if text and len(text) > 15:
                    parts.append(text)

        return "\n".join(parts)


# ---------------------------------------------------------------------------
# PDF Çıkarıcı
# ---------------------------------------------------------------------------

class PDFExtractor:
    def extract(self, pdf_bytes: bytes, url: str) -> RawDocument | None:
        # pdfplumber daha iyi tablo/sütun desteği sağlar
        try:
            import pdfplumber  # type: ignore[import]
            return self._pdfplumber(pdf_bytes, url)
        except ImportError:
            pass
        except Exception:
            pass

        # Yedek: pypdf
        try:
            return self._pypdf(pdf_bytes, url)
        except Exception:
            return None

    # ── pdfplumber yolu ──────────────────────────────────────────────────

    def _pdfplumber(self, pdf_bytes: bytes, url: str) -> RawDocument | None:
        import pdfplumber  # type: ignore[import]

        pages_text: list[str] = []
        meta_title = ""

        with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
            raw_meta = pdf.metadata or {}
            meta_title = (raw_meta.get("Title") or "").strip()

            for i, page in enumerate(pdf.pages, 1):
                parts: list[str] = []

                # Tablo çıkarma
                tables = page.extract_tables() or []
                for table in tables:
                    if not table:
                        continue
                    md = _plumber_table_to_md(table)
                    if md:
                        parts.append(md)

                # Metin çıkarma
                text = (page.extract_text() or "").strip()
                if text:
                    parts.insert(0, text)

                if parts:
                    pages_text.append(f"[Sayfa {i}]\n" + "\n\n".join(parts))

        if not pages_text:
            return None

        title = meta_title or _title_from_url(url)
        source = urlparse(url).netloc or "belek.edu.tr"
        return RawDocument(
            url=url,
            title=title,
            body="\n\n".join(pages_text),
            fmt="pdf",
            metadata={"source": source, "num_pages": len(pages_text)},
        )

    # ── pypdf yolu (yedek) ───────────────────────────────────────────────

    def _pypdf(self, pdf_bytes: bytes, url: str) -> RawDocument | None:
        from pypdf import PdfReader  # type: ignore[import]

        reader = PdfReader(io.BytesIO(pdf_bytes))
        meta   = reader.metadata or {}
        title  = (meta.get("/Title") or "").strip()

        pages_text: list[str] = []
        for i, page in enumerate(reader.pages, 1):
            text = (page.extract_text() or "").strip()
            if text and len(text) > 30:
                pages_text.append(f"[Sayfa {i}]\n{text}")

        if not pages_text:
            return None

        title = title or _title_from_url(url)
        source = urlparse(url).netloc or "belek.edu.tr"
        return RawDocument(
            url=url,
            title=title,
            body="\n\n".join(pages_text),
            fmt="pdf",
            metadata={"source": source, "num_pages": len(pages_text)},
        )


# ---------------------------------------------------------------------------
# DOCX Çıkarıcı
# ---------------------------------------------------------------------------

class DOCXExtractor:
    def extract(self, docx_bytes: bytes, url: str) -> RawDocument | None:
        try:
            from docx import Document as DocxDoc  # type: ignore[import]
        except ImportError:
            return None

        try:
            doc    = DocxDoc(io.BytesIO(docx_bytes))
            props  = doc.core_properties
            title  = (props.title or "").strip()
            parts: list[str] = []

            for para in doc.paragraphs:
                text  = para.text.strip()
                style = para.style.name if para.style else ""
                if not text:
                    continue
                if "Heading 1" in style:
                    parts.append(f"\n## {text}\n")
                elif "Heading 2" in style or "Heading 3" in style:
                    parts.append(f"\n### {text}\n")
                else:
                    parts.append(text)

            for table in doc.tables:
                md_rows: list[str] = []
                for i, row in enumerate(table.rows):
                    cells = [c.text.strip() for c in row.cells]
                    md_rows.append("| " + " | ".join(cells) + " |")
                    if i == 0:
                        md_rows.append("|" + " :--- |" * len(cells))
                if md_rows:
                    parts.append("\n".join(md_rows))

            if not parts:
                return None

            title  = title or _title_from_url(url)
            source = urlparse(url).netloc or "belek.edu.tr"
            return RawDocument(
                url=url,
                title=title,
                body="\n".join(parts),
                fmt="docx",
                metadata={"source": source},
            )
        except Exception:
            return None


# ---------------------------------------------------------------------------
# Yardımcı: tablo → Markdown
# ---------------------------------------------------------------------------

def _table_to_markdown(table_el) -> str:
    """BeautifulSoup <table> etiketini Markdown tabloya dönüştür."""
    rows = table_el.find_all("tr")
    if not rows:
        return ""

    result: list[str] = []
    for i, row in enumerate(rows):
        cells = row.find_all(["th", "td"])
        if not cells:
            continue
        texts = [c.get_text(separator=" ", strip=True) for c in cells]
        result.append("| " + " | ".join(texts) + " |")
        if i == 0:
            result.append("|" + " :--- |" * len(cells))

    return "\n".join(result) if len(result) > 1 else ""


def _plumber_table_to_md(table: list[list]) -> str:
    """pdfplumber tablo listesini Markdown tabloya dönüştür."""
    result: list[str] = []
    for i, row in enumerate(table):
        if not row:
            continue
        cells = [str(cell or "").replace("\n", " ").strip() for cell in row]
        result.append("| " + " | ".join(cells) + " |")
        if i == 0:
            result.append("|" + " :--- |" * len(cells))
    return "\n".join(result) if len(result) > 1 else ""


def _title_from_url(url: str) -> str:
    """URL'den insanca okunabilir başlık üret."""
    name = url.rstrip("/").split("/")[-1]
    name = re.sub(r"\.(pdf|docx|html?)$", "", name, flags=re.I)
    name = re.sub(r"[_\-]", " ", name)
    return name.strip()


# ---------------------------------------------------------------------------
# Dispatcher: format tespiti + yönlendirme
# ---------------------------------------------------------------------------

_HTML_EXT = frozenset({".html", ".htm", ".php", ".asp", ".aspx"})
_PDF_EXT  = frozenset({".pdf"})
_DOCX_EXT = frozenset({".docx"})


def dispatch(url: str, content_type: str, raw: bytes) -> RawDocument | None:
    """
    Gelen URL + içerik tipine göre doğru çıkarıcıyı seç.

    content_type: 'html' | 'pdf' | 'docx' | '' (tahmin edilecek)
    raw         : ham bytes (HTML için UTF-8 decode edilecek)
    """
    ext = _url_ext(url)

    # PDF
    if content_type == "pdf" or ext in _PDF_EXT or "pdf" in content_type:
        return PDFExtractor().extract(raw, url)

    # DOCX
    if content_type == "docx" or ext in _DOCX_EXT or "word" in content_type:
        return DOCXExtractor().extract(raw, url)

    # HTML (varsayılan)
    if content_type == "html" or ext in _HTML_EXT or "html" in content_type or not ext:
        html = raw.decode("utf-8", errors="replace")
        return HTMLExtractor().extract(html, url)

    return None


def _url_ext(url: str) -> str:
    path = urlparse(url).path.lower()
    dot  = path.rfind(".")
    return path[dot:] if dot != -1 else ""
